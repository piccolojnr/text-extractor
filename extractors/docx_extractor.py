from docx import Document
from PIL import Image
import pytesseract
import io


def extract_text_from_docx(file_path):
    """Extract text from a DOCX."""
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX: {e}")
    return text


def extract_text_and_images_from_docx(file_path):
    """Extract text and perform OCR on images in a DOCX."""
    text = ""
    try:
        document = Document(file_path)

        # Extract text from paragraphs
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"

        # Extract images from the document
        for rel in document.part.rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                image_stream = io.BytesIO(image_data)
                pil_image = Image.open(image_stream)

                # Perform OCR on the image
                ocr_text = pytesseract.image_to_string(pil_image)
                text += f"\nOCR from Embedded Image:\n{ocr_text}\n"
    except Exception as e:
        raise RuntimeError(f"Failed to extract text and images from DOCX: {e}")
    return text
